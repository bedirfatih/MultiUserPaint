#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MultiUserPaint - Kod Analiz ve Dokümantasyon Üreti
Tüm class'ları mantıksal sıraya göre düzenleyip DOCX dosyası oluşturur.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_documentation():
    doc = Document()
    
    # Başlık
    title = doc.add_heading('MultiUserPaint - Kod Analizi', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Tüm Class\'ları Açıklayan Teknik Dokümantasyon')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(12)
    subtitle_format.font.italic = True
    
    info = doc.add_paragraph(f'Oluşturma Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Boş satır
    
    # RabbitMQ Avantajları
    doc.add_heading('RabbitMQ Tercihinin Nedenleri ve Avantajları', 1)
    
    doc.add_heading('1. RabbitMQ Nedir ve Neden Kullanıldı?', 2)
    doc.add_paragraph(
        'Arkadaşlar, bu projemizde biz mesaj kuyruk sistemi olarak RabbitMQ kullanmayı tercih ettik. '
        'RabbitMQ, AMQP protokolüne dayalı açık kaynak bir mesaj aracısı (message broker). '
        'Basitçe söylemek gerekirse, istemciler ile sunucu arasında veri gönderirken, doğrudan bağlantı kurmak yerine, '
        'bir aracı aracılığıyla iletişim kuruyor. Bu aracı, mesajları saklıyor, yönetiyor ve doğru yerlere gönderiyor.'
    )
    
    doc.add_heading('2. RabbitMQ\'nun Temel Avantajları', 2)
    
    doc.add_heading('A) Asenkron İletişim (Asynchronous Communication)', 3)
    doc.add_paragraph(
        'RabbitMQ ile sunucu mesajları hemen göndermek zorunda değil. Mesaj kuyruğa konuyor ve '
        'alıcı hazır olduğunda alıyor. Bu sayede sunucu istemciyi beklemek zorunda kalmaz. '
        'Tek yönlü haberleşme gibi düşünebilirsiniz: birisi mektup yıkıyor, öteki hazır olduğunda okuyor.'
    )
    
    doc.add_heading('B) Fanout Exchange (Yayın-Abone Modeli)', 3)
    doc.add_paragraph(
        'RabbitMQ\'de fanout exchange kullanarak bir mesajı birden çok alıcıya aynı anda gönderebiliyoruz. '
        'Örneğin: bir kullanıcı sisteme giriş yaptığında, o andan itibaren sistemde olan HERKES bunu öğreniyor. '
        'TCP\'de bunu yapmak için sunucu tüm istemcileri tek tek kontrol edip mesaj göndermek zorunda. '
        'RabbitMQ\'de ise exchange mesajı otomatik olarak tüm abone olan kuyruklara tutuyor.'
    )
    
    doc.add_heading('C) Dosya-Spesifik Broadcasting (paint.file.{fileId})', 3)
    doc.add_paragraph(
        'Bir dosya üzerinde çalışan istemcilere sadece o dosyayla ilgili değişiklikleri gönderebiliyoruz. '
        'Diyelim 50 tane dosya açık var. Dosya 3 değiştiğinde, sadece dosya 3\'ü açan istemciler haberdar oluyor. '
        'Diğer 49 dosyayı açan istemciler hiçbir şey almıyor. Bu çok daha verimli bir sistemdir.'
    )
    
    doc.add_heading('D) Bağlantı Bağımsızlığı (Connection Independence)', 3)
    doc.add_paragraph(
        'TCP ile sunucu ve istemci arasında direkt bir bağlantı olması gerekiyor. '
        'İstemci offline olursa mesaj kayboluyor. RabbitMQ\'de ise mesajlar kuyruklarda bekliyebiliyor. '
        'İstemci geri geldiğinde kuyruklanan tüm mesajları alabilir. Ayrıca sunucu ve istemci arasında firewall olsa bile, '
        'ikisi de RabbitMQ aracısına bağlanabilirse iletişim kurabiliyor.'
    )
    
    doc.add_heading('E) Ölçeklenebilirlik (Scalability)', 3)
    doc.add_paragraph(
        'TCP ile her yeni istemci için sunucuda yeni bir thread veya bağlantı açması gerekiyor. '
        'Binlerce istemci olduğunda sunucu çok yüklenebilir. RabbitMQ ise bu yükü kendi üzerine alıyor. '
        'Sunucu sadece RabbitMQ\'ya bağlanıyor, istemciler de bağlanıyor. RabbitMQ merkezi yöneticisi rolünü oynuyor.'
    )
    
    doc.add_heading('F) Garantili Teslimat (Delivery Guarantee)', 3)
    doc.add_paragraph(
        'RabbitMQ, mesajın alıcıya ulaştığını teyit ediyor. TCP\'de birisi mesajı kaçırırsa yeniden gönderme işini '
        'programcının yazması gerekiyor. RabbitMQ bunu otomatik yapar. Mesaj alındıktan sonra onay (ACK) gelirse '
        'RabbitMQ mesajı siliyor, yoksa tekrar gönderebiliyor.'
    )
    
    doc.add_heading('3. TCP vs RabbitMQ Karşılaştırması', 2)
    
    # Tablo verilerini hazırlıyorum
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Özellik'
    hdr_cells[1].text = 'TCP (Direkt Socket)'
    hdr_cells[2].text = 'RabbitMQ (AMQP)'
    
    rows_data = [
        ['Bağlantı Modeli', 'Direkt sunucu-istemci', 'Merkezi aracı (Broker)'],
        ['İletişim Türü', 'Senkron (blocking)', 'Asenkron (non-blocking)'],
        ['Broadcast', 'Manuel olarak tüm istemcilere gönder', 'Fanout exchange otomatik'],
        ['Bağlantı Kesintisi', 'Mesaj kaybı', 'Kuyrukta bekler, sonra teslim'],
        ['Sunucu Yükü', 'Her istemci için resource', 'Merkezi yönetim, az yük'],
        ['Ölçeklenebilirlik', 'Sınırlı (thread havuzu)', 'Yüksek (binlerce istemci)'],
        ['Firewall/NAT', 'Sorun olabilir', 'Her iki taraf aracıya bağlanırsa çalışır'],
        ['Teslimat Garantisi', 'Yok, timeout ile kontrol', 'Evet, ACK mekanizması'],
        ['Seçici Teslimat', 'Zor (manuel filter)', 'Kolay (Exchange + Binding)'],
        ['Yapılandırma', 'Basit ama sınırlı', 'Karmaşık ama çok esnek']
    ]
    
    for row_data in rows_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row_data[0]
        row_cells[1].text = row_data[1]
        row_cells[2].text = row_data[2]
    
    doc.add_heading('4. Bu Projede Neden RabbitMQ Tercih Ettik?', 2)
    doc.add_paragraph(
        'Bu MultiUserPaint uygulamasında çok sayıda istemcinin aynı anda paylaşılan dosyalar üzerinde çalışması gerekiyor. '
        'Birisi çizim yaparken diğerleri bunu anında görmek zorunda. TCP ile bunu yapmak mümkün ama çok karmaşık. '
        'Çünkü her çizim olayında sunucu tüm ilişkili istemcileri bulup ilk mesaj göndermek zorunda. '
        'RabbitMQ ile ise: bir dosya opened ise, biz o dosyaya ait bir exchange kuruyoruz. '
        'O dosyayı açan tüm istemciler otomatik olarak o exchange\'e abone oluyor. '
        'Çizim mesajı geldiğinde, RabbitMQ otomatik olarak tüm abone olanlara dağıtıyor. Çok daha temiz ve güvenilir.'
    )
    
    doc.add_heading('5. RabbitMQ Architecture Bu Projede', 2)
    p = doc.add_paragraph()
    p.add_run('paint.server - ').bold = True
    p.add_run('Tüm istemcilerin sunucuya gönderdiği mesajlar için ortak kuyruk\n')
    p.add_run('paint.global - ').bold = True
    p.add_run('Tüm istemcilere ulaşması gereken mesajlar (USER_JOIN, FILE_LIST)\n')
    p.add_run('paint.file.{fileId} - ').bold = True
    p.add_run('Spesifik bir dosyayı açan istemcilere giden mesajlar (DRAW_BROADCAST)\n')
    p.add_run('paint.client.{id} - ').bold = True
    p.add_run('Her istemcinin özel yanıt kuyruğu (point-to-point mesajlar)')
    
    doc.add_paragraph(
        'Bu yapı sayesinde ağ trafiği çok azalıyor. Doğru mesaj doğru yerlere gidiyor. '
        'Sunucu hiçbir filtreleme yapmak zorunda kalmıyor, RabbitMQ bunu halletmiş. '
        'Kodumuz daha temiz, daha bakımlanabilir ve ölçeklenebilir oluyor.'
    )
    
    doc.add_page_break()
    toc_items = [
        ('I. Protocol & Ortak Katman (Foundation)', [
            '1. ProtocolConstants',
            '2. MessageType',
            '3. FSMState',
            '4. Message',
            '5. MessageEncoder',
            '6. MessageDecoder',
            '7. MQConfig'
        ]),
        ('II. Sunucu Katmanı (Server)', [
            '8. PaintServer',
            '9. NIOSelector',
            '10. ClientSession',
            '11. MessageDispatcher',
            '12. MessageFramer',
            '13. LoginHandler',
            '14. FileHandler',
            '15. DrawHandler',
            '16. ClipboardHandler',
            '17. FileStore',
            '18. FileMetadata',
            '19. SessionRegistry',
            '20. AutoSaveScheduler',
            '21. MQBroker',
            '22. MQServerTransport'
        ]),
        ('III. İstemci Katmanı (Client)', [
            '23. PaintClient',
            '24. ConnectionManager',
            '25. NetworkReader',
            '26. NetworkWriter',
            '27. MQNetworkReader',
            '28. MQNetworkWriter',
            '29. MainFrame',
            '30. CanvasPanel',
            '31. ToolbarPanel',
            '32. FileListPanel',
            '33. LoginDialog'
        ])
    ]
    
    for category, items in toc_items:
        doc.add_heading(category, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # PART I: Protocol & Ortak Katman
    doc.add_heading('BÖLÜM I: Protocol & Ortak Katman (Foundation)', 1)
    doc.add_paragraph(
        'Yazılımın temel altyapısını oluşturan sınıflar. İstemci ve sunucu arasında veri iletişimini '
        'standardize eden protokol tanımları, mesaj formatları ve durum yönetimi burada yer alır.'
    )
    
    # 1. ProtocolConstants
    doc.add_heading('1. ProtocolConstants', 2)
    doc.add_paragraph(
        'Burada, arkadaşlar, biz bir iletişim protokolünün temel kurallarını tanımlıyoruz. '
        'Aynı bir oyunun kuralları gibi düşünebilirsiniz. Herkes aynı kuralı bilirse oyun oynayabilir. '
        'Bu sınıf da istemci ve sunucunun "ortak dili" gibi çalışıyor.'
    )
    doc.add_heading('Temel Kurallar:', 3)
    doc.add_paragraph(
        'Önce sihirli sayı diyoruz (0xCAFE). Bir mesaj geldi mi, biz ilk bakıyoruz: "Hey, sihirli sayı doğru mu?" '
        'Doğruysa, bu ağımızın mesajı demek. Yanlışsa, hata var veya başka bir protokolden veri geldi. '
        'Sürüm numarası (VERSION = 0x01) ise gelecekte protokolü güncelledimiz zaman ne yapacağımızı söylüyor. '
        '"Bu mesaj 1 numaralı protokol ile yazıldı" gibi.'
    )
    doc.add_paragraph(
        'Başlık boyutu (HEADER_SIZE = 8 byte) sabit. Bu 8 byte içinde sihirli sayı, sürüm, mesaj türü ve '
        'payload uzunluğu yazılı oluyor. Böylece sunucu ilk 8 byte\'ı okuyunca hangi veriyi alacağını biliyor.'
    )
    doc.add_paragraph(
        'Maksimum payload (10 MB) sınırı koyuyoruz. Bir resim dosyası (800×600×4 byte = ~2 MB) gelse bile, '
        'birisi bir terabaylık dosya göndermeye çalışmadan koruma altına alıyoruz. Sunucunun hafızası sınırsız değil çünkü.'
    )
    doc.add_paragraph(
        'Port numarası, buffer boyutu, kullanıcı adı karakter sınırı, auto-save aralığı, keepalive timeout... '
        'Hepsi burada merkezi olarak tutuluyoruz. Eğer biz buradaki 60 sayısını 120 yaparsak, '
        'sistem her 120 saniyede dosyaları kaydeder. Başka hiçbir yerde değişiklik yapmamız gerekmez. '
        'Bu da kodun bakımını çok kolaylaştırıyor.'
    )
    
    # 2. MessageType
    doc.add_heading('2. MessageType', 2)
    doc.add_paragraph(
        'Şimdi şu mesaj türlerini tanımlıyoruz. Sistemi düşünün: kullanıcı giriş yapabilir, '
        'dosya oluşturabilir, çizim yapabilir, pano işlemi yapabilir... Her aksiyon farklı bir mesaj türü. '
        'Biz tüm bu türleri enum olarak tanımlıyoruz. Her birine benzersiz bir heksadesimal kod veriyoruz.'
    )
    doc.add_heading('Mesaj Türleri Kategorileri:', 3)
    doc.add_paragraph(
        'Giriş mesajları (0x01, 0x02, 0x03): Kullanıcı sisteme giriş yapıyor, başarılı mı başarısız mı.'
    )
    doc.add_paragraph(
        'Dosya mesajları (0x10-0x1A): Dosya oluştur, aç, kaydet, sil, listesi iste... Tüm dosya işlemleri.'
    )
    doc.add_paragraph(
        'Çizim mesajları (0x20, 0x21): Kullanıcı fare ile çizim yapıyor, ben bunu sunucuya "Hey, '
        'bu dosyada şu noktadan şu noktaya çizgi çiz!" diye mesaj gönderiyorum.'
    )
    doc.add_paragraph(
        'Pano mesajları (0x30-0x33): Copy, Cut, Paste. Çok istemcili ortamda birinin kestiği şey '
        'diğerinin bilgisayarında paste edilebiliyor.'
    )
    doc.add_paragraph(
        'Kullanıcı mesajları (0x40-0x43): Birisi sisteme giriş yaptı, çıktı, '
        'tüm istemcilere "Hey, şu kişi sistemde" diye haber veriyoruz.'
    )
    doc.add_paragraph(
        'Canvas anlık görüntü (0x60-0x62): Yeni birisi bir dosyayı açtı. '
        'Dosya zaten açık olan birisi var mı? Varsa onun canvas\'ından şu anın resmi gönder, '
        'disk\'ten okuma. Bu sayede herkesin aynı görüntüye bakması garantileniyor.'
    )
    doc.add_paragraph(
        'Keepalive (0xF0, 0xF1): Ping-pong. Aradan uzun süre ses çıkmazsa sunucu "Hey, yaşıyor musun?" diye ping yolluyor. '
        'İstemci pong ile cevap veriyor. "Evet yaşıyorum, endişelenme" gibi.'
    )
    doc.add_paragraph(
        'Hata mesajı (0xFF): Birisi yanlış bir işlem yaptı, sunucu hata mesajı gönderiyor.'
    )
    
    # 3. FSMState
    doc.add_heading('3. FSMState', 2)
    doc.add_paragraph(
        'Sonlu Durum Makinesi diyoruz buna. Tarihiyle basit: bir oyuncu düşünün. '
        'Oyunun başında bekleme durumunda. Menu açılıyor, oyunu seçiyor, oyuna giriyor. '
        'Oyun sırasında birisi oyunu durduruyor, tekrar başlıyor... Bu durumlar FSM durumları. '
        'Sunucuda da her bağlanan istemci bu durumlardan birinde oluyor.'
    )
    doc.add_heading('Sunucu Tarafında Durumlar:', 3)
    doc.add_paragraph(
        'HANDSHAKE: İstemci bağlandı ama henüz giriş yapmadı. Sunucu "Bana kim olduğunu söyle" diyor. '
        'Bu durumda sadece LOGIN_REQ mesajı kabul ediliyor. Başka mesaj gelirse, sunucu "Önce giriş yap!" diye hata veriyor.'
    )
    doc.add_paragraph(
        'ACTIVE: Giriş başarılı. Şimdi kullanıcı dosya açabilir, çizim yapabilir, pano işlemi yapabilir. '
        'Tüm normal operasyonlar bu durumda gerçekleşiyor.'
    )
    doc.add_paragraph(
        'SAVING: Dosya diske yazılıyor. Sunucu bu sırada yoğun. '
        'Mesajlar gelse bile, ping-pong dışında işlenmiyor. Disk yazması bittikten sonra tekrar ACTIVE oluyor.'
    )
    doc.add_paragraph(
        'CLOSING: Güzelce çıkıyor. Tüm kaynakları temizliyor, dosyaları kapatıyor. Sorun yok, sana da hoşça kal diyecek.'
    )
    doc.add_paragraph(
        'ERROR: Aaaa bir hata oldu! Düzeltilemiyor. Bu oturum artık işe yaramıyor, kapatılacak. '
        'İstemci yeniden bağlanması gerekiyor.'
    )
    
    doc.add_heading('İstemci Tarafında Durumlar:', 3)
    doc.add_paragraph(
        'DISCONNECTED: Sunucuya bağlı değil. İstemci uygulaması açık olsa da, hiçbir verisi sunucuda yok.'
    )
    doc.add_paragraph(
        'CONNECTING: "Hey sunucu, bana açık mı?" Socket açıyor, bağlantı kuruyor.'
    )
    doc.add_paragraph(
        'LOGGING_IN: "Benim adım şu, al giriş yap beni" LOGIN_REQ göndermiş, cevap bekliyor.'
    )
    doc.add_paragraph(
        'CONNECTED: "Tamamdır, artık ben de senin gibi aktif!" Giriş başarılı, dosyaları görüyor, çizim yapabiliyor.'
    )
    doc.add_paragraph(
        'RECONNECTING: İstemci düştü ama sunucu henüz kapamadı. Tekrar bağlanmaya çalışıyor. '
        'Bazen bağlantı kesintileri olur (internet problemi), otomatik yeniden bağlanılıyor.'
    )
    
    # 4. Message
    doc.add_heading('4. Message', 2)
    doc.add_paragraph(
        'Bu çok basit bir sınıf. Bir mesajın en temel şeklidir. İçinde sadece iki şey var: '
        'mesajın türü (MessageType) ve taşıdığı veriler (byte dizisi). '
        'Tamamen değişmez (immutable), bir kere oluşturulduktan sonra değiştirilmiyor. '
        'Güvenli ve basit.'
    )
    doc.add_paragraph(
        'Kullanımı çok basit: Message msg = Message.of(MessageType.LOGIN_OK, payload); '
        'Ve o kadar. Artık bu message istediğimiz yere gönderilebiliyor.'
    )
    
    # 5. MessageEncoder
    doc.add_heading('5. MessageEncoder', 2)
    doc.add_paragraph(
        'Bu sınıf yazılı verileri ağ üzerinden gönderilebilecek byte dizilerine dönüştürüyor. '
        'Biz programcılar Java nesneleriyle rahat çalışıyoruz (String, int, byte[]). '
        'Fakat ağ üzerinden sadece byte gönderilmek zorunda. İşte buraya MessageEncoder devreye giriyor.'
    )
    doc.add_paragraph(
        'Örneğin: Kullanıcı "Ali" isminde giriş yapmak istiyor. Biz bunu şöyle yapıyoruz:'
    )
    doc.add_paragraph(
        'byte[] payload = MessageEncoder.encodeLoginReq("Ali");'
    )
    doc.add_paragraph(
        'MessageEncoder ne yapıyor? "Ali" stringini UTF-8\'e dönüştürüyor, önüne de uzunluğunu (2 byte) yazıyor. '
        'Sonra bunu bir Message\'a sarıyor, ağ formatını ('
    )
    p = doc.add_paragraph()
    p.add_run('MAGIC, VER, TYPE, LENGTH, PAYLOAD')
    p.add_run(') ekliyor ve son olarak ağ üzerinden gönderilebilecek byte dizisini veriyor.')
    
    # 6. MessageDecoder
    doc.add_heading('6. MessageDecoder', 2)
    doc.add_paragraph(
        'Bu, MessageEncoder\'ın tam tersi. Ağdan gelen byte dizilerini (ki bunlar garip, '
        'başında sihirli sayı falan var) tekrar insanın anlayabileceği şekle çeviriyoruz.'
    )
    doc.add_paragraph(
        'Örneğin sunucu MESSAGE geldi, ağdan. MessageDecoder.decodeLoginReq() çağırıyoruz. '
        'O da byte dizisini parse ediyor, sihirli sayıyı ve başlığı çıkartıyor, '
        'geriye kalan payload\'ı alıyor, onu da UTF-8\'den string\'e çeviriyor. '
        'Ve sonuç: LoginReqPayload nesnesi, içinde username = "Ali" yazıyor. Çok rahat!'
    )
    doc.add_paragraph(
        'Decoder içinde LoginReqPayload, LoginOkPayload, FileCreateReqPayload gibi '
        'payload türleri var. Bunlar basit POJO sınıfları (Plain Old Java Objects), '
        'sadece verileri tutuyorlar. Message türüne göre doğru payload sınıfı kullanılıyor.'
    )
    
    # 7. MQConfig
    doc.add_heading('7. MQConfig', 2)
    doc.add_paragraph(
        'RabbitMQ kullandığımızda, mesajların nasıl yönetileceğini tanımlıyoruz. '
        'Bu konfigürasyondaki "topoloji" demek, kuyruklar ve exchange\'ler nasıl bağlanacak?'
    )
    doc.add_heading('RabbitMQ Topoloji Detaylı Anlatım:', 3)
    doc.add_paragraph(
        'Şöyle düşünün: Bir postane hayal edin. Postane birden çok kutusu var. '
        'Sunucudan gelen tüm mesajlar paint.server kutusuna gidiyor. '
        'RabbitMQ bu kutuyu okuyuyor, mesajı alıyor ve şuna şu kutuya koy diyerek dağıtıyor.'
    )
    doc.add_paragraph(
        'paint.global: Tüm istemcilere ulaşması gereken mesajlar. Birisi sisteme giriş yaptı (USER_JOIN), '
        'herkes bunu bilmeli. Bunlar paint.global fanout exchange\'e gidiyor. '
        'Fanout exchange ne yapıyor? "Bunu bana bağlı olan herkese gönder" diyor. '
        'Tüm istemcilerin kuyrukları buna bağlı olduğu için, herkes mesajı alıyor.'
    )
    doc.add_paragraph(
        'paint.file.{fileId}: Spesifik bir dosya için exchange. Dosya 5 açılmış olsun. '
        'Dosya 5 üzerinde çizim yaparsan, DRAW_BROADCAST paint.file.5 exchange\'ine gidiyor. '
        'Paint.file.5\'e abone olanlar (sadece dosya 5\'i açanlar) mesajı alıyor. '
        'Diğer dosyaları açanlar hiçbir şey almıyor. Verimli, değil mi?'
    )
    doc.add_paragraph(
        'paint.client.{id}: Her istemcinin özel kutusu. Benim client ID\'im 12345. '
        'Sunucu bana mesaj göndermek isterse, paint.client.12345\'e koyuyor. '
        'Sadece ben (client 12345) onun abonesi, sadece ben alıyorum. Point-to-point iletişim bu.'
    )
    doc.add_paragraph(
        'Konfigürasyon dosyası RabbitMQ\'nun bağlantı detaylarını da tutuyor: '
        'Host (localhost), port (5672), kullanıcı adı (guest), şifre (guest)... '
        'Bu da merkezi bir yerde tutulup, başka yerden değişiklik yapılmıyor.'
    )
    
    doc.add_page_break()
    
    # PART II: Sunucu Katmanı
    doc.add_heading('BÖLÜM II: Sunucu Katmanı (Server)', 1)
    doc.add_paragraph(
        'Şimdi arkadaşlar, sunucu tarafına geçiyoruz. Sunucu tüm istemcilerin bağlantılarını yönetiyor, '
        'paylaşılan dosyaları diske kaydediyor ve mesajları yönlendiriyor. '
        'TCP ile çalışabilir, ya da RabbitMQ ile çalışabilir. İkisi de tamamen işliyor.'
    )
    
    # 8. PaintServer
    doc.add_heading('8. PaintServer', 2)
    doc.add_paragraph(
        'Bu tüm sunucunun başı. Uygulamayı çalıştırdığında en önce bu sınıf çalışıyor. '
        'Java main() metodu burada. Ne yapıyor? Bütün sistemi başlatıyor, kullandığı nesneleri oluşturuyor.'
    )
    doc.add_heading('Kullanım:', 3)
    doc.add_paragraph(
        'TCP modunda: java PaintServer 9090 /path/to/files'
    )
    doc.add_paragraph(
        'RabbitMQ modunda: java PaintServer --mq rabbitmq.example.com /path/to/files'
    )
    
    doc.add_heading('Başlatılan Altyapı:', 3)
    doc.add_paragraph(
        'SessionRegistry: "Sistem, kim bağlı?" sorusunun cevabı. Tüm istemcileri burada tutuyor.'
    )
    doc.add_paragraph(
        'FileStore: Disk\'te dosyaları okup yazıyor. 800×600×4 byte = 1.92 MB resim.'
    )
    doc.add_paragraph(
        'DiskPool: 2 thread havuzu. Disk çok yavaş olabilir, bu yüzden bu işleri ayrı threaddlerde yapıyoruz. '
        'Ana sunucu thread\'i bloke olmaz.'
    )
    doc.add_paragraph(
        'MessageDispatcher: Mesajlar geldiğinde, bunlar kimin işine giriyorsa ona yönlendirir. '
        'Giriş mesajı → LoginHandler, Çizim mesajı → DrawHandler, vb.'
    )
    doc.add_paragraph(
        'LoginHandler, FileHandler, DrawHandler, ClipboardHandler: Her biri farklı türde mesajları işleyen uzmanlar.'
    )
    doc.add_paragraph(
        'AutoSaveScheduler: Her 60 saniyede bir tüm açık dosyaları diske kaydediyor. Veri kaybı olmasın diye.'
    )
    doc.add_paragraph(
        'NIOSelector (TCP) veya MQServerTransport (RabbitMQ): Ağ iletişimini yönetiyorlar. '
        'Hangisini seçtiysen ona göre başlatılıyor.'
    )
    
    # 9. NIOSelector
    doc.add_heading('9. NIOSelector', 2)
    doc.add_paragraph(
        'Bu TCP modundaki kalp atışı. NIO demek Non-Blocking I/O. Bunun neden önemli olduğunu anlatayım. '
        'Normal socket programming\'de, her istemci için bir thread açarsın. 1000 istemci = 1000 thread. '
        'Her thread hafıza tutuyor, context switch çok. Sunucu çöp olur.'
    )
    doc.add_paragraph(
        'NIO ile? Tek bir thread tüm 1000 istemciyi yönetiyor. Nasıl? Selector diye bir şey var. '
        'Selector diyor: "Ey tüm socketler, hangisinin veri gelmişse bana söyle." '
        'Bir socket "Benim veri var" derse, selector gider o socket\'ten okur. '
        'Başkası "Benimde var" derse, ondan da okur. Sırayla herkese bakar. '
        'Bu çok verimli ve hızlı.'
    )
    doc.add_heading('Yapıyor Olduğu İşler:', 3)
    doc.add_paragraph(
        'Accept: Yeni istemci mi geliyor? (Mesela port 9090\'a bağlanmaya çalışıyor) '
        'Selector kabul ediyor, o istemci için ClientSession oluşturuyor, registry\'ye ekliyor.'
    )
    doc.add_paragraph(
        'Read: Hangi istemcilerin veri gelmişse onlardan okuyor. Okunan byte\'ları MessageFramer\'a besliyor. '
        'MessageFramer mesaj birleştiriyor ve "Bir tam mesaj hazır!" dediğinde, '
        'Dispatcher\'a gönderiyor işlenmek üzere.'
    )
    doc.add_paragraph(
        'Write: İstemciye yanıt yazılacak mı? (writeQueue\'de mesaj var mı?) '
        'Varsa, session\'ın writeQueue\'sinden alıp istemciye gönderiyor.'
    )
    doc.add_paragraph(
        'Keepalive: Aradan 90 saniye veri almadığı istemci var mı? Evet, öldürmüş kabul ediyoruz. '
        'Bağlantı kapanıyor. Ping-pong mekanizması de var aslında, ama çok uzun sürerse bu çalışıyor.'
    )
    
    # 10. ClientSession
    doc.add_heading('10. ClientSession', 2)
    doc.add_paragraph(
        'Bağlı her istemci için sunucuda bir ClientSession nesnesi var. '
        'Bu obje, "Ben kim im? Hangi dosyaları açtım? Panom\'da ne var? Son ping zamanı ne zaman?" '
        'gibi tüm detaylarını tutuyor. Bir istemci kaç yaşındaki çocuksa, ClientSession onun "özet bilgisi".'
    )
    doc.add_heading('İçerdiği Ana Veriler:', 3)
    doc.add_paragraph(
        'sessionId: Her istemci için benzersiz numara. "Sen 12345 numara mı? O zaman seni tanıyorum."'
    )
    doc.add_paragraph(
        'channel (TCP modunda): SocketChannel. Doğrudan haberleşme tüpü gibi. RabbitMQ\'de null.'
    )
    doc.add_paragraph(
        'state: FSMState. Şu anda HANDSHAKE mı, ACTIVE mi, SAVING mi durumda?'
    )
    doc.add_paragraph(
        'username: Giriş yaptığı kullanıcı adı. "Ali" giriş yaptıysa, username = "Ali".'
    )
    doc.add_paragraph(
        'Mesaj çerçeveleme tamponları (readBuffer, headerBuf, payloadBuf): '
        'Gelen veriler burada birleştiriliyor. Örneğin, 5 bayt geldiyse burada tutulur, '
        'sonra 3 bayt daha gelse, toplam 8 bayt oluyor ve başlık tamam diyelim.'
    )
    doc.add_paragraph(
        'writeQueue: Gönderilecek mesajlar kuyruk. SessionRegistry mesajı buraya koyuyor, '
        'NIOSelector\'ın write loop\'u buradan alıp istemciye gönderiyor.'
    )
    doc.add_paragraph(
        'openFileIds: Şu istemci hangi dosyaları açmış? Set olarak tutuyor. '
        'Dosya 1, 3, 5 açıksa, {1, 3, 5} seti tutuyor.'
    )
    doc.add_paragraph(
        'clipboardData: Kopya-yapıştır işlemleri. Kullanıcı Ctrl+C yaptıysa, '
        'o bölgenin pixel verisi burada saklanıyor. Ctrl+V yaptığında, buradan çekiliyor.'
    )
    doc.add_paragraph(
        'RabbitMQ modundaki ekstra veriler: mqReplyTo (istemcinin özel kuyruk adı), mqSink (gönderme işlevi).'
    )
    
    # 11. MessageDispatcher
    doc.add_heading('11. MessageDispatcher', 2)
    doc.add_paragraph(
        'Mesaj yönlendiricisi. Bir mesaj geldiğinde, "Bu neyin işi? Kimin yapması lazım?" diye soruyor. '
        'Giriş mesajı geldi mi → LoginHandler\'a git. Çizim mesajı geldi mi → DrawHandler\'a git. '
        'Dosya mesajı geldi mi → FileHandler\'a git. Trafik polisi gibi.'
    )
    doc.add_paragraph(
        'Ama bir kural var: istemcinin durumu önemli. HANDSHAKE\'de sadece LOGIN_REQ kabul ediliyor. '
        'Başka mesaj gelirse, "Hey! Önce kendini tanıt!" hata veriyor. '
        'ACTIVE\'de, tüm normal mesajlar işleniyor. SAVING\'de, disk yazması bitmeden mesajlar (ping hariç) işlenmiyor. '
        'ERROR\'de, kapanış başlatılıyor.'
    )
    
    # 12. MessageFramer
    doc.add_heading('12. MessageFramer', 2)
    doc.add_paragraph(
        'Ağdan gelen ham byte akışını tamamlanmış mesajlara dönüştüren sihircisi. '
        'Şöyle bir sorun var: Network sabit paketler göndermiyor. '
        'Kümlaştırabilir. Bölebilir. Mesaj 1000 byte\'sa, 500 geliyor, sonra 250, sonra 250. '
        'Biz bunları birleştirmek zorundayız.'
    )
    doc.add_paragraph(
        'MessageFramer iki aşamada çalışıyor:'
    )
    doc.add_paragraph(
        'Aşama 1 - Başlık: İlk 8 byte için bekliyoruz. Sihirli sayı kontrol, mesaj türü, payload boyutunu öğreniyoruz.'
    )
    doc.add_paragraph(
        'Aşama 2 - Payload: Artık boyutunu biliyoruz. Mesela 2000 byte payload var diyordu başlık. '
        '200 byte geldiyse, 1800 daha bekliyoruz. Tamamdığında, Message nesnesi oluşturup dispatcher\'a veriyoruz.'
    )
    doc.add_paragraph(
        'Böylece 10 MB\'lık bir resim bile sorun olmaz. 64 KB puffer ile bile, büyük mesajları işleyebiliriz.'
    )
    
    # 13. LoginHandler
    doc.add_heading('13. LoginHandler', 2)
    doc.add_paragraph(
        'Kullanıcı giriş yaptığında çalışan handler. "Benim adım Ali, beni sisteme al!" dediğinde, '
        'LoginHandler devreye giriyor.'
    )
    doc.add_heading('Yaptığı Kontroller:', 3)
    doc.add_paragraph(
        'Kullanıcı adı geçerli mi? 3-16 karakter, sadece harf-rakam-underscore mi? '
        'Değilse, hatamı söylüyoruz. "Adında özel karakter var, kullanma."'
    )
    doc.add_paragraph(
        'Başka birisi bu adı kullanıyor mu? Evet ise, "Maalesef adını başkası kullanıyor, başka seç" diyoruz.'
    )
    doc.add_paragraph(
        'Geçti her kontrol? OK! "Giriş başarılı, session ID = 12345 senin" diyoruz. '
        'state\'i ACTIVE yapıyoruz. SessionRegistry\'ye ekliyoruz. '
        'Ve tüm BAĞLI istemcilere "HEY! Ali sisteme girdi!" mesajı yayınlıyoruz. '
        'Herkes bunu alınca bilir ki Ali online\'da.'
    )
    
    # 14. FileHandler
    doc.add_heading('14. FileHandler', 2)
    doc.add_paragraph(
        'Dosya işlemleri için sorumlu handler. En karmaşık handler bu. Neden? Çünkü dosya açıp kapatma, '
        'bir dosyayı birden çok kişi aynı anda görmesi, canvas senkronizasyonu, hepsi burada. '
        'Ayrıca disk operasyonları yavaş olduğu için, bunu ayrı thread\'lerde yapıyor (DiskPool).'
    )
    doc.add_heading('Dosya Oluştur (handleCreate):', 3)
    doc.add_paragraph(
        'Kullanıcı "Yeni Dosya - 800x600" dedi. FileHandler bunu DiskPool\'a gönderiyor. '
        'DiskPool thread\'i: beyaz 800x600 resim yaratıyor, PNG olarak diske yazıyor, metadata oluşturuyor. '
        'Bittiğinde, istemciye "OK tamamdır, dosya ID = 5" yanıtı gidiyor. '
        'Tüm istemcilere de "Yeni dosya var!" haber veriliyor.'
    )
    doc.add_heading('Dosya Aç (handleOpen):', 3)
    doc.add_paragraph(
        'Kullanıcı var olan bir dosyayı açıyor. Fakat burada ince bir yer var. '
        'Dosya 5 zaten açık olsa da başka birisi tarafından, '
        'canvas diske yazıldığından beri değişmiş olabilir. O kişinin en son görüntüsünü almalıyız. '
        'Yoksa diskteki eski versiyon gösterir. Yanlış olur.'
    )
    doc.add_paragraph(
        'Bunun için, FileHandler kontrol ediyor: dosyayı açan başka biri var mı? '
        'Varsa, ona "Hey, sen bu dosyayı açmışsın, şu an canvas durumunu gönder" diyor (CANVAS_SNAPSHOT_REQ). '
        'O kişi snapshot gönderiyor. Snapshot geliyor, yeni açan kişiye veriliyor. '
        'Eğer kimse açmamışsa, disk\'ten okuyuyor.'
    )
    doc.add_heading('Dosya Kaydet (handleSave):', 3)
    doc.add_paragraph(
        'Kullanıcı Ctrl+S yaptı. Harita veri sunucuya geliyor. FileHandler bunu DiskPool\'a gönderiyor. '
        'DiskPool PNG olarak yazıyor. Bittiğinde state SAVING\'den ACTIVE\'e dönüyor.'
    )
    doc.add_heading('Dosya Listesi (handleList):', 3)
    doc.add_paragraph(
        'Kullanıcı "Bana tüm dosyaları göster" dedi. FileHandler FileStore\'dan tüm metadata alıyor, '
        'bunları paketliyor ve TÜMAPTA istemcilere (broadcast) gönderiyor. '
        'Herkes dosya listesini güncelliyor.'
    )
    
    # 15. DrawHandler
    doc.add_heading('15. DrawHandler', 2)
    doc.add_paragraph(
        'Çizim handler. En basit handler. Neden? Çünkü sadece istediği dosyayı açan '
        'diğer istemcilere mesajı yayınlıyor.'
    )
    doc.add_paragraph(
        'Kullanıcı A dosya 3\'de, noktadan (100,100) noktaya (200,200) mavi çizgi çiziyor. '
        'DRAW_EVENT mesajı geliyor. DrawHandler: Dosya 3, araç=LINE, renk=blue, koordinatlar, '
        'sen (A) kimin olduğunu söyle (A username\'i ekle) ve bunu DRAW_BROADCAST olarak '
        'dosya 3\'ü açan herkese gönder. '
        'Dosya 3\'ü açan B, C, D istemcileri mesajı alıyor, kendi canvaslarında çizgi görüyorlar. '
        'Gerçek zamanlı işbirliği!'
    )
    
    # 16. ClipboardHandler
    doc.add_heading('16. ClipboardHandler', 2)
    doc.add_paragraph(
        'Kesme-yapıştır işlemleri. Kullanıcı bir bölge seçip Ctrl+C yapıyor. '
        'ClientSession\'da clipboardData kaydediliyor. Dosya 3\'ü açan başkası Ctrl+V yapıyor. '
        'Pano verisi o noktaya yapıştırılıyor. '
        'Aynı dosyayı açan herkes bunu görebiliyor. Paylaşılan pano.'
    )
    
    # 17. FileStore
    doc.add_heading('17. FileStore', 2)
    doc.add_paragraph(
        'Tüm dosyaları disk\'te yöneten depo. PNG formatında kaydediyor, pixel verilerini ARGB byte dizilerine çeviriyor. '
        'Database gibi düşünebilirsiniz ama file-based.'
    )
    doc.add_paragraph(
        'Her dosyanın bir ID\'si var: 1, 2, 3... Disk\'te: file_1.png, file_2.png, file_3.png olarak kaydediliyor. '
        'Metadata de tutuluyor: dosyanın sahibi, boyutu, son değiştirilme zamanı...'
    )
    doc.add_paragraph(
        'Pixel format: 4 byte per pixel (ARGB). Alfa = opaklık, RGB = renk. '
        '800x600 resim = 800×600×4 = 1,920,000 byte = 1.9 MB. Bu kadar.'
    )
    
    # 18. FileMetadata
    doc.add_heading('18. FileMetadata', 2)
    doc.add_paragraph(
        'Bir dosyanın tüm bilgilerini tutan hafif veri taşıyıcısı. Adı ne? Sahibi kim? Kaç byte? Ne zaman değişti?'
    )
    doc.add_paragraph(
        'İçinde: fileId, filename, owner, width, height, lastModified. Hepsi bu. '
        'Dosya listesi gönderirken, FileStore tüm dosyaların metadata\'sını gönderir, pixel verilerini değil. '
        'Böylece çok az data transfer oluyor.'
    )
    
    # 19. SessionRegistry
    doc.add_heading('19. SessionRegistry', 2)
    doc.add_paragraph(
        'Sistem\'in "telefonlar rehberi". Şu an kimler online? Kimlerin açık dosyaları var? '
        'Bir dosyayı kim açmış? Hepsi burada. Çok önemli bir sınıf.'
    )
    doc.add_heading('Yapısı:', 3)
    doc.add_paragraph(
        'Harita 1 - Session ID ile arama: "Session 12345 kim?" Hemen bulur.'
    )
    doc.add_paragraph(
        'Harita 2 - Username ile arama: "Ali sistem\'de mi?" Hemen bulur. '
        'Böylece başkası "Ali" adında giriş yapamaz (double login engellenir).'
    )
    doc.add_heading('Broadcast İşlevleri:', 3)
    doc.add_paragraph(
        'broadcastToAll(mesaj): Tüm istemcilere "Birisi sisteme girdi" mesajı yayınla. '
        'Exclude kimliğini kullanabilir - "Ben\'e gönderme, ben zaten biliyorum" gibi.'
    )
    doc.add_paragraph(
        'broadcastToFileViewers(fileId, mesaj): Dosya 3\'ü açan tüm istemcilere mesaj gönder. '
        'Öteki dosyaları açan istemciler ilgilenmez.'
    )
    doc.add_paragraph(
        'Eğer RabbitMQ modunda ise, bu fonksiyonlar broker üzerinden çalışıyor (fanout exchange). '
        'Eğer TCP modunda ise, tüm istemcileri manual olarak gezdip gönderiliyor.'
    )
    
    # 20. AutoSaveScheduler
    doc.add_heading('20. AutoSaveScheduler', 2)
    doc.add_paragraph(
        'Sistem "kendi kendini kaydeden" mekanizması. Her 60 saniyede bir alarm çalıyor, '
        'tüm açık dosyaları diske kaydediyor. Böylece sunucu çöpse bile, son 60 saniyenin verisi kaybı en çok olur.'
    )
    doc.add_paragraph(
        'Nasıl çalışıyor? SessionRegistry\'den tüm açık dosya ID\'lerini toplayıyor. '
        'Her dosya için FileStore\'dan pixel verisi okuyuyor (disk\'ten)... bekle, disk\'ten mi? '
        'Evet, canvas her istemcide tutuluyor (backBuffer). Dosya açık istemcilerden snapshot alması gerekirse, '
        'snapshot yollanıyor, ama auto-save sadece diskteki versiyon saklanıyor. Yani bilinen versiyon.'
    )
    doc.add_paragraph(
        'Hata olursa, hata log\'u tutuluyor ama sistem çalışmaya devam ediyor.'
    )
    
    # 21. MQBroker
    doc.add_heading('21. MQBroker', 2)
    doc.add_paragraph(
        'RabbitMQ aracısıyla iletişim kuran sunucu tarafı adapter. '
        'Bağlantı açıyor, kuyrukları ve exchange\'leri deklarasyon ediyor, mesaj yayınlıyor.'
    )
    doc.add_heading('Yapılandırma:', 3)
    doc.add_paragraph(
        'Host: rabbitmq.example.com (veya localhost). Port: 5672 (AMQP default portu). '
        'VirtualHost, kullanıcı, şifre: rabbitmq konfigürasyonunda belirlenen.'
    )
    doc.add_heading('Topoloji Kurulumu:', 3)
    doc.add_paragraph(
        'Paint.server kuyruğunu deklarasyon ediyor: tüm istemcilerden gelen mesajlar buraya giriyor.'
    )
    doc.add_paragraph(
        'Paint.global fanout exchange: USER_JOIN, USER_LEAVE, FILE_LIST_RESP, bu exchange\'e gidiyor. '
        'Hepsi broadcast türü mesaj.'
    )
    doc.add_paragraph(
        'Paint.file.{fileId} exchange\'leri lazy decleration: Dosya 5 üzerinde çizim ilk yapılırken, '
        'exchange oluşturuluyor. Dosya kapatılırsa exchange silinebiliyor (auto-delete).'
    )
    doc.add_heading('Mesaj Gönderme:', 3)
    doc.add_paragraph(
        'sendToClient(replyTo, frame): Point-to-point, bir istemciye doğrudan mesaj.'
    )
    doc.add_paragraph(
        'broadcastGlobal(frame): Fanout exchange\'e mesaj, herkese geliyor.'
    )
    doc.add_paragraph(
        'broadcastToFile(fileId, frame): Dosya exchange\'ine mesaj, o dosyayı açanlar alıyor.'
    )
    
    # 22. MQServerTransport
    doc.add_heading('22. MQServerTransport', 2)
    doc.add_paragraph(
        'RabbitMQ modunda NIOSelector yerine kullanılan consumer. '
        'Paint.server kuyruğundan mesajları okuyor, işliyor, sonuç gönderiyor.'
    )
    doc.add_heading('İşlem Akışı:', 3)
    doc.add_paragraph(
        'basicConsume(paint.server): Paint.server kuyruğundan mesaj almaya başla. Autoack=false, manual ACK ile.'
    )
    doc.add_paragraph(
        'Mesaj geldiğinde, body byte dizisi, AMQP headers\'tan sessionId ve replyTo çıkart.'
    )
    doc.add_paragraph(
        'ClientSession bul, yoksa oluştur. MessageDispatcher\'a göster. Sonuç döndür. BasicAck ile onayda.'
    )
    doc.add_paragraph(
        'Keepalive sweeper: Ayrı thread, 30 saniye interval\'de kontrol. '
        'Eğer bir session 90 saniye ses çıkmazsa, close flag atıyor, kapatılıyor.'
    )
    
    doc.add_page_break()
    
    # PART III: İstemci Katmanı
    doc.add_heading('BÖLÜM III: İstemci Katmanı (Client)', 1)
    doc.add_paragraph(
        'Şimdi arkadaşlar, istemci tarafına geçiyoruz. Burada kullanıcı ile doğrudan temas kurulacak yer. '
        'Swing GUI (click, drag, menu) ve ağ iletişiminin beraber çalıştığı katman.'
    )
    
    # 23. PaintClient
    doc.add_heading('23. PaintClient', 2)
    doc.add_paragraph(
        'Tüm istemci uygulamasının başlama noktası. Main metodu burada. '
        'Ne yapıyor? Swing\'in look-and-feel\'ini sistem temasına ayarlıyor, '
        'EDT (Event Dispatch Thread) üzerinde MainFrame oluşturuyor, '
        'login dialog\'u açıyor. O kadar basit.'
    )
    
    # 24. ConnectionManager
    doc.add_heading('24. ConnectionManager', 2)
    doc.add_paragraph(
        'İstemcinin "kalp ve beyni". Bağlantı yaşam döngüsünü ve ağ iletişimini yönetiyor. '
        'Çok önemli olan: TCP ve RabbitMQ modlarını kapsüllediği için, '
        'GUI katmanı hangisini kullandığını bilmez. İkiside aynı API kullanıyor.'
    )
    doc.add_heading('TCP Modu Başlatma:', 3)
    doc.add_paragraph(
        'connect("localhost", 9090) → Socket aç, sunucuya bağlan. '
        'NetworkReader ve NetworkWriter thread\'leri başlat.'
    )
    doc.add_heading('RabbitMQ Modu Başlatma:', 3)
    doc.add_paragraph(
        'connectMQ("localhost", 5672) → RabbitMQ aracısına bağlan. '
        'Client-side özel kuyruğu oluştur, subscribe et. '
        'MQNetworkReader ve MQNetworkWriter başlat.'
    )
    doc.add_heading('Sağlanan Metotlar:', 3)
    doc.add_paragraph(
        'sendLoginReq(username): Giriş mesajı gönder (MessageEncoder kullanarak)'
    )
    doc.add_paragraph(
        'sendFileCreateReq, sendFileOpenReq, sendFileSaveReq, ...: Dosya işlemleri'
    )
    doc.add_paragraph(
        'sendDrawEvent: Çizim olayı (araç, renk, koordinatlar)'
    )
    doc.add_paragraph(
        'sendClipboardCopy, sendClipboardCut, sendClipboardPasteReq: Pano işlemleri'
    )
    doc.add_paragraph(
        'disconnect(): Bağlantı kapat, cleanup yap'
    )
    
    # 25. NetworkReader
    doc.add_heading('25. NetworkReader', 2)
    doc.add_paragraph(
        'TCP modundaki mesaj alıcı. Ayrı daemon thread\'te çalışır. '
        'Blocking InputStream\'den okur. Bir mesaj gelene kadar bekler.'
    )
    doc.add_paragraph(
        'Mesaj geldiğinde? Başlık (8 byte) okur, magic kontrol eder, '
        'payload boyutunu öğrenir. Payload (payloadLen byte) okur. '
        'Message nesnesi oluştur, SwingUtilities.invokeLater() ile EDT\'de onMessage callback\'i çağır. '
        'GUI thread\'inde gösterebilmesi için EDT\'ye koyuyor.'
    )
    doc.add_paragraph(
        'Bağlantı kesintisi varsa (EOFException), onDisconnect callback\'i çağırıyor. '
        'GUI state\'i "Disconnected"\'e alıyor.'
    )
    
    # 26. NetworkWriter
    doc.add_heading('26. NetworkWriter', 2)
    doc.add_paragraph(
        'TCP modundaki mesaj gönderici. BlockingQueue kullanarak thread-safe çalışıyor. '
        'EDT\'den send(byte[]) çağrısı gelirse, veriler queue\'ye konuyor. '
        'Daemon writer thread\'i queue\'den alıp OutputStream\'e yazıyor. '
        'Block olmadığı için, GUI responsive kalıyor.'
    )
    
    # 27. MQNetworkReader
    doc.add_heading('27. MQNetworkReader', 2)
    doc.add_paragraph(
        'RabbitMQ modundaki mesaj alıcı. '
        'BasicConsumer kullanarak client kuyruğundan mesaj tüketiyor. '
        'Login olmadan: client exclusive queue + global fanout exchange\'e subscribe. '
        'LOGIN_OK alındıktan sonra: sessionId header\'ı sağlanıyor. '
        'Dosya açılınca: o dosya exchange\'ine subscribe (notifyFileOpened). '
        'Dosya kapatılınca: unsubscribe (notifyFileClosed).'
    )
    
    # 28. MQNetworkWriter
    doc.add_heading('28. MQNetworkWriter', 2)
    doc.add_paragraph(
        'RabbitMQ modundaki mesaj gönderici. NetworkWriter gibi çalışıyor. '
        'BlockingQueue\'den frame alıp, paint.server kuyruğuna yayınlıyor. '
        'Header\'lara: replyTo (client kuyruğu adı) ve sessionId (login\'den sonra) ekleniyor. '
        'RabbitMQ otomatik yönlendirme yapıyor.'
    )
    
    # 29. MainFrame
    doc.add_heading('29. MainFrame', 2)
    doc.add_paragraph(
        'Ana uygulama penceresi. Bütün GUI öğelerinin toplandığı yer. '
        'Ekranın solunda dosya listesi, üstünde toolbar, ortasında tabbed pane (açık canvaslar), '
        'altında status bar. Bütün bunun koordinatörü MainFrame.'
    )
    doc.add_heading('Yapısı:', 3)
    doc.add_paragraph(
        'ToolbarPanel: Araç seçme, renk seçme, stroke kalınlığı seçme. "Pencil mi, Line mi?" burada.'
    )
    doc.add_paragraph(
        'FileListPanel: Sol sidebar\'da tüm dosyaları listele. "Dosya 1 (Ali), Dosya 2 (Ayşe)" şeklinde. '
        'Double-click → o dosya açılıyor, yeni tab oluşuyor.'
    )
    doc.add_paragraph(
        'JTabbedPane: Her açılan dosya için bir tab ve CanvasPanel. Dosya 1 tab\'ında çiziyorken, '
        'Dosya 2 tab\'ına geçebiliyoruz. Current toolbar ayarları (renk, araç) tablar arasında paylaşılıyor.'
    )
    doc.add_paragraph(
        'Menu: Dosya (Yeni, Kaydet, Yenile), Edit (Copy, Cut, Paste), Server (Bağlan, Ayırlan).'
    )
    doc.add_paragraph(
        'Status Bar: "Bağlantı durumu: Çevrimiçi", "File saved", vb. mesajlar.'
    )
    doc.add_heading('Mesaj İşleme:', 3)
    doc.add_paragraph(
        'onMessage() metodu tüm sunucu mesajlarını alıyor. Gelen mesajın türüne göre: '
        'LOGIN_OK mi? Durumu güncelle. FILE_LIST_RESP mi? FileListPanel\'i güncelle. '
        'DRAW_BROADCAST mi? Aktif canvas\'da çizgi göster. USER_JOIN mi? Status bar\'da "Ali katıldı" yaz.'
    )
    
    # 30. CanvasPanel
    doc.add_heading('30. CanvasPanel', 2)
    doc.add_paragraph(
        'Gerçek çizim yüzeyimiz. 800x600 resolution. JPanel\'den extends ediyor, '
        'mouse event\'lerini yakalayıp, çizim işlemleri yönetiyor.'
    )
    doc.add_heading('Teknik Detaylar:', 3)
    doc.add_paragraph(
        'Dual-buffer rendering: backBuffer (kalıcı, sunucudan alınan versiyon) ve '
        'overlayBuffer (geçici, şu anda çizilen stroke). '
        'Kullanıcı çizim yaparken, overlayBuffer\'a çiziliyor (lokal feedback). '
        'Mouse release\'de stroke gönderiliyor sunucuya. '
        'Sunucudan confirmation gelince (veya diğer kullanıcılardan DRAW_BROADCAST), '
        'backBuffer\'a merge ediliyor. Çok realist görünüyor.'
    )
    doc.add_heading('Araçlar:', 3)
    doc.add_paragraph('Pencil: Serbest çizim', style='List Bullet')
    doc.add_paragraph('Line: Düz çizgi', style='List Bullet')
    doc.add_paragraph('Rect: Dikdörtgen', style='List Bullet')
    doc.add_paragraph('Ellipse: Elips', style='List Bullet')
    doc.add_paragraph('Fill: Renk doldur (bucket)', style='List Bullet')
    doc.add_paragraph('Eraser: Silgi (beyaz pixel yaz)', style='List Bullet')
    doc.add_paragraph('Selection: Ctrl+C/X/V için seçim alanı', style='List Bullet')
    
    # 31. ToolbarPanel
    doc.add_heading('31. ToolbarPanel', 2)
    doc.add_paragraph(
        'Basit toolbar. 6 araç butonu, renk seçici, stroke kalınlığı combobox. '
        'Kullanıcı burada araç seçiyorsa, CanvasPanel\'a "Şimdi Pencil\'i kullan" diye haber veriliyor.'
    )
    
    # 32. FileListPanel
    doc.add_heading('32. FileListPanel', 2)
    doc.add_paragraph(
        'Sol sidebar. JList kullanıyor. Sunucudan FILE_LIST_RESP gelince, '
        'liste güncelleniyor. "dosya_adı (sahibi)" formatında gösteriliyor. '
        'Double-click → FileHandler.handleOpen() çağrılıyor, dosya açılıyor.'
    )
    
    # 33. LoginDialog
    doc.add_heading('33. LoginDialog', 2)
    doc.add_paragraph(
        'İlk açıldığında gösterilen modal dialog. '
        'Mode seçimi (Socket/RabbitMQ), sunucu host, port, username alanları. '
        'Mode değiştiğinde port otomatik değişiyor (9090 → 5672 veya tersi). '
        'Connect butonu clicked → ConnectionManager.connect() veya connectMQ() çağrılıyor.'
    )
    doc.add_paragraph(
        'Başarılı bağlantıda dialog kapanıyor, MainFrame görünüyor. '
        'Hata varsa, error popup gösteriliyor, tekrar deneyebilmesi için dialog açık kalıyor.'
    )
    
    doc.add_page_break()
    
    # Sonuç
    doc.add_heading('Sistem Genel Yapısı ve İş Akışı', 1)
    
    doc.add_heading('Mimarinin Katmanları:', 2)
    doc.add_paragraph(
        'Arkadaşlar, bu sistem üç ana katmandan oluşuyor. '
        'En üstte GUI (Swing), ortada bağlantı (ConnectionManager ve Network), '
        'en altta protokol ve ağ (Wire Format + Transport).'
    )
    p = doc.add_paragraph()
    p.add_run('Presentation Layer (Swing)\n').bold = True
    p.add_run('MainFrame, CanvasPanel, ToolbarPanel, FileListPanel, LoginDialog\n\n')
    p.add_run('Connection Abstraction (ConnectionManager)\n').bold = True
    p.add_run('TCP ve RabbitMQ modlarını şeffaflaştırır\n\n')
    p.add_run('Transport Layer\n').bold = True
    p.add_run('TCP: NetworkReader/Writer + NIOSelector\n')
    p.add_run('RabbitMQ: MQNetworkReader/Writer + MQServerTransport\n\n')
    p.add_run('Protocol Layer\n').bold = True
    p.add_run('MessageEncoder/Decoder, MessageFramer\n\n')
    p.add_run('Foundation Layer\n').bold = True
    p.add_run('MessageType, FSMState, ProtocolConstants, Message')
    
    doc.add_heading('Gerçek Zamanlı Çizim Örneği Detaylı Akış:', 2)
    doc.add_paragraph(
        'Ali dosya 3\'de resim yapıyor. Noktadan (100,100) (200,200) mavi çizgi çizmek istiyoruz.'
    )
    steps = [
        'Ali mouse\'u (100,100)\'de basıyor → CanvasPanel.onPressed() çağrılıyor',
        'Ali (200,200)\'ye sürüklüyor → overlayBuffer\'a çizgi gösteriliyor (local feedback)',
        'Ali mouse\'u bırakıyor → onReleased() çağrılıyor',
        'CanvasPanel DRAW_EVENT mesajı oluşturuyor: fileId=3, tool=LINE, color=blue, x1=100, y1=100, x2=200, y2=200',
        'ConnectionManager.sendDrawEvent() çağrılıyor',
        'TCP: NetworkWriter queue\'ye mesajı konuyor, daemon thread gönder',
        'RabbitMQ: MQNetworkWriter paint.server kuyruğuna publish ediyor',
        'Sunucu mesajı alıyor → NIOSelector (TCP) veya MQServerTransport (RabbitMQ)',
        'MessageDispatcher DRAW_EVENT\'i DrawHandler\'a yönlendidir',
        'DrawHandler: "Dosya 3 açan kimler var?" diye SessionRegistry\'ye sorur',
        'Dosya 3\'ü açan tüm istemciler bulunur (Ali, Bob, Carol olsun)',
        'DRAW_BROADCAST mesajı bu üçüne gönderilir (Ali\'ye de)',
        'Her istemcinin NetworkReader/MQNetworkReader mesajı alıyor',
        'MainFrame.onMessage() DRAW_BROADCAST işliyor',
        'CanvasPanel.onDrawBroadcast() backBuffer\'a çizgiyi merge ediyor',
        'paint() metodu çağrılıyor, çizgi ekranda görünüyor',
        'Ali, Bob ve Carol aynı anda çizgiyi görüyorlar!'
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(step, style='List Number')
    
    doc.add_heading('Başlıca Tasarım Kararları ve Neden Önemli:', 2)
    doc.add_paragraph(
        'Dual-Transport Architecture: Biz bir proje yaptık ama iki moda uyuyor. '
        'Hocamız TCP ister mi, RabbitMQ ister mi, ikisinde de çalışıyor. '
        'ConnectionManager abstraction\'ı sayesinde GUI hiçbir şey değiştirmeden çalışıyor.'
    )
    doc.add_paragraph(
        'Non-Blocking Server (NIO): Binlerce istemciyi tek thread\'le yönetebiliyoruz. '
        'Eğer her istemci için thread açsaydık, server ölürdü. '
        'NIO sayesinde scalable ve hızlı.'
    )
    doc.add_paragraph(
        'Asenkron Disk I/O: Dosya okuma-yazma yavaş. DiskPool kullanarak, '
        'bu işler ayrı threaddlerde yapılıyor. Sunucu bloklamamış oluyor. '
        'Mesajlar gelmeye devam ediyor.'
    )
    doc.add_paragraph(
        'FSM-based Dispatching: Her client\'in durumuna göre mesajlar işleniyor. '
        'Hata handling çok temiz ve güvenilir oluyor. '
        'State transition\'lar açık ve anlaşılır.'
    )
    doc.add_paragraph(
        'Dual-Buffer Rendering: Canvas\'ta çizim responsiv gözüküyor. '
        'Local feedback (overlayBuffer) anında, server confirmation sonrası commit ediliyor. '
        'Bağlantı yavaşsa bile GUI freeze olmuyor.'
    )
    doc.add_paragraph(
        'RabbitMQ Fanout Exchange\'ler: Doğru mesaj doğru kişilere gidiyor. '
        'Broadcast yapılırken sunucu hiç bir şey yapmak zorunda kalmıyor. '
        'RabbitMQ hepsini halletmiş, very efficient.'
    )
    doc.add_paragraph(
        'Auto-Save Scheduler: Periyodik olarak tüm açık dosyalar kaydediliyor. '
        'Sunucu çöpse, data loss en çok 60 saniye. '
        'Güvenilir, otomatik, user action gerekmez.'
    )
    doc.add_paragraph(
        'Canvas Snapshot Senkronizasyonu: Yeni birisi bir dosyaya katılırsa, '
        'o anın canvas\'ı tutarlı hale getiriliyor. '
        'Disk versiyonu değil, live versiyonu kullanıyor. '
        'Böylece everyone on the same page.'
    )
    
    # Dosyayı kaydet
    output_path = r'c:\Users\Caner\Desktop\Drive\İUC\Ağ ve İletişim Teknolojileri\proje\MultiUserPaint\MultiUserPaint_Kod_Analizi.docx'
    doc.save(output_path)
    print(f'✅ DOCX dosyası başarıyla oluşturuldu: {output_path}')
    return output_path

if __name__ == '__main__':
    create_documentation()
